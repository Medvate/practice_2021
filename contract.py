from docx import Document
from docx.shared import Pt


PRIMARY_TITLE = ['Наименование НИОКР (шифр)',
                 'Сроки выполнения НИОКР',
                 'Стоимость НИОКР',
                 'Контракт (номер, дата)',
                 'ТТЗ (кем утверждено, дата)',
                 'Заказчик',
                 'Исполнитель НИОКР',
                 'Главный конструктор (научный руководитель) НИОКР (организация, должность, ФИО, контакты)',
                 'Соисполнитель СЧ НИОКР',
                 'НИО, осуществляющая ВНС НИОКР',
                 'ТТЗ на ВНС НИОКР (кем утверждено, дата)',
                 'Тематическая карточка и (когда и кем утверждены, с кем согласованы)',
                 'Справа-обоснование на ВНС НИОКР (когда и кем утверждены, с кем согласованы',
                 'Научный руководитель работы ВНС НИОКР (должность, звание, ФИО, контактный телефон, факс организации)',
                 'Ответственный исполнитель работы ВНС НИОКР (должность, звание, ФИО, контактный телефон)']

SUBSEQUENT_TITLES = ['Отчетный период (начало (ден.месяц.год – окончание (ден.месяц.год))',
                     'Наименование НИОКР (шифр)',
                     'Этап выполнения НИОКР',
                     '',
                     '',
                     'Причины отставания (при наличии)',
                     'Характеристика этапа',
                     'Работы, выполняемые в отчетном периоде',
                     'Перечень утвержденных и согласованных документов в отчетном периоде',
                     'Перечень неисполненных документов (причины)',
                     'Выездной контроль',
                     'Наименование организации',
                     '',
                     'Общие проблемные вопросы по выполнению НИОКР',
                     'Меры, принятые по исключению срыва выполнения НИОКР',
                     'Головным исполнителем',
                     'Заказчиком',
                     'Органом военного управления',
                     'НИО, осуществляющее ВНС НИОКР',
                     'Акт приемки выполнения этапа (работы) (дата, кем утвержден)',
                     'Присвоение литеры , , приказы о принятии на вооружение (снабжение), когда и кем утвержден, '
                     'сведения о поставках потребителям ВВСТ',
                     'Предложения по выполнению НИОКР (закрытие, перенос сроков (причины))']


class Contract:
    def __init__(self, filepath: str) -> None:
        self._docx = Document(filepath)
        self._contract_type = NotImplemented
        self._table_data = dict()

        self._find_contract_type()
        self._find_table_data()

    def _find_contract_type(self):
        type_as_str = self._docx.paragraphs[2].text
        type_as_str = type_as_str.replace(' ', '')

        if type_as_str == '(первичная)':
            self._contract_type = 'primary'
        elif type_as_str == '(последующая)':
            self._contract_type = 'subsequent'
        else:
            raise ValueError('Договор не имеет типа!')

    def _find_table_data(self):
        table = self._docx.tables[0]

        if self._contract_type == 'primary':
            self._parse_table_from_primary_doc(table)
        elif self._contract_type == 'subsequent':
            self._parse_table_from_subsequent_doc(table)
        else:
            raise ValueError('Договор не имеет таблицы!')

    def _parse_table_from_primary_doc(self, table) -> None:
        for index, row in enumerate(table.rows):
            row_title = row.cells[0].text

            if index < 15 and PRIMARY_TITLE[index] == row_title:
                self._table_data[row_title] = row.cells[1].text
            elif row_title == '':
                if row.cells[1].text == 'Наименование организации (задачи)' and \
                        row.cells[2].text == 'ТТЗ на СЧ ВНС НИОКР (когда, кем согласовано и утверждено)':
                    continue
                else:
                    raise ValueError(f'Ошибка в заголовках "Наименование организации (задачи)" и'
                                     f'"ТТЗ на СЧ ВНС НИОКР (когда, кем согласовано и утверждено)".')
            elif row_title == 'Соисполнитель ВНС НИОКР':
                new_task = {'Наименование организации (задачи)': row.cells[1].text,
                            'ТТЗ на СЧ ВНС НИОКР (когда, кем согласовано и утверждено)': row.cells[2].text}
                if self._table_data.get(row_title):
                    self._table_data[row_title] = [*self._table_data[row_title], new_task]
                else:
                    self._table_data[row_title] = [new_task]
            else:
                raise ValueError(f'Ошибка в заголовках: "{row_title}", ожидалось "{PRIMARY_TITLE[index]}".')

    def _parse_table_from_subsequent_doc(self, table) -> None:
        for index, row in enumerate(table.rows):
            row_title = row.cells[0].text

            if SUBSEQUENT_TITLES[index] == row_title:
                if index in [0, 14]:
                    continue
                elif index == 3:
                    if row.cells[1].text == 'Сроки выполнения (плановые)' and \
                            row.cells[3].text == 'Сроки выполнения (фактические)':
                        continue
                    else:
                        raise ValueError('Ошибка в заголовках "Сроки выполнения (плановые)" или'
                                         '"Сроки выполнения (фактические)".')
                elif index == 4:
                    self._table_data['Сроки выполнения (плановые)'] = row.cells[1].text
                    self._table_data['Сроки выполнения (фактические)'] = row.cells[3].text
                elif index == 11:
                    if row.cells[1].text == 'Дата, представители НИО осуществляющее ВНС' and \
                            row.cells[2].text == 'Проблемные вопросы':
                        continue
                    else:
                        raise ValueError('Ошибка в заголовках "Дата, представители НИО осуществляющее ВНС" или'
                                         '"Проблемные вопросы".')
                elif index == 12:
                    self._table_data['Дата, представители НИО осуществляющее ВНС'] = row.cells[1].text
                    self._table_data['Проблемные вопросы'] = row.cells[2].text
                else:
                    self._table_data[row_title] = row.cells[1].text
            else:
                raise ValueError(f'Ошибка в заголовках: "{row_title}", ожидалось "{PRIMARY_TITLE[index]}".')

    def is_relatives(self, contract) -> bool:
        if self._table_data['Наименование НИОКР (шифр)'] == contract._table_data['Наименование НИОКР (шифр)']:
            return True
        else:
            return False

    def merge(self, contract):
        for key, value in contract._table_data.items():
            if self._table_data.get(key):
                new_value = self._table_data[key]
                if new_value != value:
                    if type(new_value) == str:
                        new_value += f'\n{value}'
                    elif type(new_value) == list:
                        new_value += value
                    self._table_data[key] = new_value
            else:
                self._table_data[key] = value


if __name__ == '__main__':
    c_1 = Contract('primary_form_1.docx')
    c_2 = Contract('subsequent_form_1.docx')

    if c_1.is_relatives(c_2):
        c_1.merge(c_2)

    print(123)

    result = Document()
    font = result.styles['Normal'].font
    result.styles['Normal'].font.name = 'Times New Roman'
    result.styles['Normal'].font.size = Pt(14)

    table = result.add_table(rows=35, cols=2, style='Table Grid')

    for row, data in zip(table.rows, c_1._table_data.items()):
        row.cells[0].text = data[0]
        row.cells[1].text = str(data[1])

    result.save('result.docx')
