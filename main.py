from docx import Document


def get_type_of_doc(document: Document) -> str:
    third_paragraph = document.paragraphs[2].text
    third_paragraph = third_paragraph.replace(' ', '')

    if third_paragraph == '(первичная)':
        return 'primary'
    elif third_paragraph == '(последующая)':
        return 'subsequent'
    else:
        raise ValueError('Не получилось определить тип договора!')


if __name__ == '__main__':
    primary_doc = Document('primary_form_1.docx')
    type_of_primary = get_type_of_doc(primary_doc)

    subsequent_doc = Document('subsequent_form_1.docx')
    type_of_subsequent = get_type_of_doc(subsequent_doc)




